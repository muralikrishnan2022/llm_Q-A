from pathlib import Path
import pdfplumber
from langchain_ollama import OllamaLLM
from docx import Document

def extract_text_from_pdf(pdf_path):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None

def generate_mcqs(text, model):
    prompt = (
        "Generate 5 multiple-choice questions based on the following text. "
    "Each question must have 4 options (A, B, C, D) and specify the correct answer. "
    "Use this exact format:\n\n"
    "Example:\n"
    "Q1. What is the capital of France?\n"
    "A) Berlin\n"
    "B) Madrid\n"
    "C) Paris\n"
    "D) Rome\n"
    "Correct Answer: C\n\n"
    f"Text:\n{text.strip()}\n\n"
    "Now generate 5 questions in this format:"
    )

    try:
        response = model.invoke(prompt)
        return response
    except Exception as e:
        print(f"Error generating MCQs with Ollama: {e}")
        return None

def save_mcqs_to_word(mcqs, output_path):
    try:
        doc = Document()
        doc.add_heading("Multiple Choice Questions", level=1)
        for line in mcqs.split("\n"):
            if line.strip():
                doc.add_paragraph(line)
        doc.save(output_path)
        print(f"MCQs saved to: {output_path}")
    except Exception as e:
        print(f"Error saving MCQs to Word document: {e}")

def main():
    # Initialize Ollama model
    try:
        ollama_model = OllamaLLM(model="llama3")
    except Exception as e:
        print(f"Error initializing Ollama model: {e}")
        return

    # Get PDF file path
    pdf_path_input = input("Enter the PDF file path: ")
    pdf_path = Path(pdf_path_input)

    if not pdf_path.is_file():
        print("Error: File does not exist at the specified path.")
        return
    if pdf_path.suffix.lower() != '.pdf':
        print("Error: File is not a .pdf file.")
        return

    # Extract text
    print("Extracting text from PDF...")
    text = extract_text_from_pdf(pdf_path)
    if not text:
        print("Failed to extract text. Exiting.")
        return

    # Save extracted text to .txt file
    output_dir = pdf_path.parent
    text_file = output_dir / f"{pdf_path.stem}_extracted.txt"
    try:
        with open(text_file, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"Extracted text saved to: {text_file}")
    except Exception as e:
        print(f"Error saving extracted text: {e}")
        return

    # Generate MCQs with Ollama
    print("Generating MCQs with Ollama...")
    mcqs = generate_mcqs(text, model=ollama_model)
    if not mcqs:
        print("Failed to generate MCQs. Exiting.")
        return

    # Save MCQs to a .docx file
    mcq_file = output_dir / f"{pdf_path.stem}_mcqs.docx"
    save_mcqs_to_word(mcqs, mcq_file)

if __name__ == "__main__":
    main()
